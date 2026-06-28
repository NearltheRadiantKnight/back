# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs")
OUT_FILE = OUT_DIR / "毕业设计管理系统后端项目文档.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
BORDER = "C9D3DF"
WHITE = "FFFFFF"
MUTED = "666666"


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, bold=False, color="000000"):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color="000000", size=10.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, col_widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, widths)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_GRAY)
        set_cell_text(hdr.cells[i], header, bold=True, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], BLUE_GRAY)
        set_cell_text(cells[0], label, bold=True, color=INK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[1], value, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_table_width(table, list(widths))
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="D8E2EF", size="4")
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=INK)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(body)
    set_run_font(run2, size=10, color="222222")
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        set_paragraph_font(p, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_paragraph_font(p, size=13, bold=True, color=BLUE)
    else:
        set_paragraph_font(p, size=12, bold=True, color=DARK_BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="222222")
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=9.5, name="Consolas", color="222222")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("毕业设计管理系统后端项目文档")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Graduation Thesis / Design Defense Management Backend")
    set_run_font(r, size=12, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("基于 Spring Boot、MyBatis-Plus、MySQL 与 AI 评语生成的后端 API 说明")
    set_run_font(r, size=12, color="333333")

    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("项目目录", "D:\\idealc\\back"),
            ("文档类型", "项目开发文档 / 后端接口与模块说明"),
            ("适用对象", "课程设计、毕业设计答辩、项目验收、开发交接"),
            ("文档版本", "V1.0"),
            ("生成日期", "2026-06-28"),
        ],
    )

    add_callout(
        doc,
        "文档说明",
        "本文档依据当前仓库代码整理，重点描述后端业务模块、数据库表、接口、运行部署、测试用例与小组分工建议。"
        "由于仓库未提供正式说明书模板，本文档采用通用软件项目开发文档结构。"
    )
    doc.add_page_break()


def revision_and_toc(doc):
    add_heading(doc, "文档修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "说明", "编写"],
        [["V1.0", "2026-06-28", "根据当前后端项目生成完整项目文档", "Codex"]],
        [1200, 1800, 4800, 1560],
    )

    add_heading(doc, "目录", 1)
    toc_items = [
        "1. 项目概述",
        "2. 技术架构与项目结构",
        "3. 系统角色与权限设计",
        "4. 功能模块设计",
        "5. 数据库设计",
        "6. 接口设计",
        "7. 关键业务流程",
        "8. 运行部署说明",
        "9. 测试设计",
        "10. 项目分工建议",
        "11. 风险问题与优化建议",
        "附录：代码依据与文件映射",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()


def section_project_overview(doc):
    add_heading(doc, "1. 项目概述", 1)
    add_body(
        doc,
        "本项目是一个面向本科毕业论文、毕业设计答辩流程的后端管理系统，主要为前端提供用户登录、院系管理、教师管理、学生管理、答辩年份和答辩小组管理、评分统计、Word 模板管理、签名上传以及 AI 评语生成等接口。"
    )
    add_body(
        doc,
        "系统围绕“院系管理员组织答辩、教师参与评分、学生完成答辩材料归档”的业务流程设计，后端采用 Spring Boot 分层架构实现，数据库使用 MySQL 存储用户、院系、学生、教师、分组、成绩与模板信息。"
    )

    add_heading(doc, "1.1 建设目标", 2)
    add_bullets(
        doc,
        [
            "统一管理学院、院系管理员、教师和学生基础数据，减少线下 Excel 或纸质表格维护成本。",
            "支持按答辩年份创建答辩小组，分配组长、教师成员和学生。",
            "支持论文与设计两类考核评分项，记录教师评分明细并计算成绩。",
            "支持模板上传、占位符校验、日期配置和签名上传，为后续生成答辩相关文档提供基础。",
            "支持调用通义千问 DashScope 接口生成论文或设计答辩评语。",
        ],
    )

    add_heading(doc, "1.2 项目范围", 2)
    add_kv_table(
        doc,
        [
            ("包含范围", "后端 API、数据库初始化脚本、文件上传配置、AI 配置文件、Word 模板相关逻辑。"),
            ("不包含范围", "当前仓库不包含完整前端页面代码，也未包含正式测试报告、接口文档平台或自动化 CI 配置。"),
            ("交付形式", "Spring Boot 后端服务，默认端口 8081；可选 Node.js 模板服务，默认端口 3000。"),
        ],
    )

    add_heading(doc, "1.3 系统核心流程", 2)
    add_numbered(
        doc,
        [
            "超级管理员维护院系和院系管理员账号。",
            "院系管理员维护教师、学生和答辩年份信息。",
            "院系管理员创建答辩小组，设置答辩组长并添加教师成员。",
            "学生被分配到答辩小组，同时记录论文或设计类型、题目等信息。",
            "教师登录系统后按年份参与答辩评分，系统保存评分明细并计算小组成绩。",
            "大组成绩录入后，系统根据每组第一名成绩计算调节系数。",
            "模板、日期、签名和 AI 评语模块为答辩材料归档提供辅助。"
        ],
    )


def section_architecture(doc):
    add_heading(doc, "2. 技术架构与项目结构", 1)
    add_heading(doc, "2.1 技术栈", 2)
    add_table(
        doc,
        ["类别", "技术/依赖", "说明"],
        [
            ("后端框架", "Spring Boot 4.0.0", "主后端服务，提供 REST API。"),
            ("语言与运行环境", "Java 17", "pom.xml 中配置 Java 17。"),
            ("ORM / 持久层", "MyBatis-Plus 3.5.15", "Mapper 层负责 SQL 查询与数据写入。"),
            ("数据库", "MySQL", "application.properties 默认连接 manager 数据库。"),
            ("文件与 Word 处理", "Apache POI 5.2.0", "用于 Word 模板读取和占位符处理。"),
            ("JSON 处理", "Jackson Databind", "处理配置文件、评分明细 JSON 等。"),
            ("AI 能力", "DashScope SDK Java", "调用通义千问模型生成答辩评语。"),
            ("可选辅助服务", "Node.js + Express", "server.js 提供早期模板管理接口，建议作为历史/辅助实现看待。"),
        ],
        [1700, 2600, 5060],
    )

    add_heading(doc, "2.2 分层架构", 2)
    add_body(doc, "当前 Java 后端采用典型 Controller、Service、Mapper、Entity 分层。Controller 接收 HTTP 请求，Service 处理业务规则，Mapper 与数据库交互，Entity / DTO 承载业务数据。")
    add_table(
        doc,
        ["层次", "目录", "职责"],
        [
            ("Controller", "src/main/java/com/world/back/controller", "定义接口路径、请求参数、响应封装。"),
            ("Service", "src/main/java/com/world/back/service", "定义业务接口。"),
            ("ServiceImpl", "src/main/java/com/world/back/serviceImpl", "实现业务逻辑，例如评分、模板、用户管理。"),
            ("Mapper", "src/main/java/com/world/back/mapper", "执行数据库查询和更新。"),
            ("Entity", "src/main/java/com/world/back/entity", "定义学生、教师、模板、成绩等实体。"),
            ("Config", "src/main/java/com/world/back/config", "Web 跨域、模板上传配置。"),
            ("Utils", "src/main/java/com/world/back/utils", "登录、路径、文件、Word 解析等工具类。"),
            ("Resources", "src/main/resources", "配置文件、AI 提示词、数据库初始化脚本。"),
        ],
        [1500, 3700, 4160],
    )

    add_heading(doc, "2.3 项目目录说明", 2)
    add_code(doc, "D:\\idealc\\back")
    add_code(doc, "├─ pom.xml                              # Java 后端依赖与构建配置")
    add_code(doc, "├─ server.js                            # 可选 Node.js 模板管理服务")
    add_code(doc, "├─ src/main/java/com/world/back         # Spring Boot 主代码")
    add_code(doc, "├─ src/main/resources/application.properties")
    add_code(doc, "├─ src/main/resources/config            # AI 配置与提示词")
    add_code(doc, "└─ src/main/resources/sql/init.sql      # MySQL 初始化脚本")

    add_heading(doc, "2.4 运行配置", 2)
    add_kv_table(
        doc,
        [
            ("服务端口", "Spring Boot 默认端口为 8081。"),
            ("数据库", "MySQL，默认库名 manager。"),
            ("模板上传目录", "./uploads/templates/。"),
            ("签名上传目录", "./uploads/signatures/。"),
            ("静态资源访问", "classpath:/static/ 与 file:./uploads/。"),
        ],
    )


def section_roles(doc):
    add_heading(doc, "3. 系统角色与权限设计", 1)
    add_body(doc, "数据库 user 表通过 role 字段区分角色。当前代码中主要角色包括超级管理员、院系管理员和教师，其中教师在部分流程中可进一步成为答辩组长。")
    add_table(
        doc,
        ["角色", "role 值", "主要权限"],
        [
            ("超级管理员", "0", "维护院系、院系管理员，查看全局基础数据。"),
            ("院系管理员", "1", "管理本院系教师、学生、答辩年份、小组、模板与统计数据。"),
            ("教师", "2", "按答辩年份登录，查看指导学生或答辩小组，参与评分。"),
            ("答辩组长", "教师扩展身份", "在小组内承担组长职责，可参与小组管理和签名归档相关流程。"),
        ],
        [1800, 1300, 6260],
    )

    add_heading(doc, "3.1 登录方式", 2)
    add_bullets(
        doc,
        [
            "普通登录接口用于管理员和院系管理员身份验证。",
            "带年份登录接口用于教师选择答辩年份后进入系统。",
            "教师是否拥有某年份权限，依据指导关系、答辩组关系或答辩组长身份判断。",
        ],
    )

    add_heading(doc, "3.2 权限设计建议", 2)
    add_body(
        doc,
        "当前代码主要依赖前端传参和业务查询进行权限区分，未看到完整 JWT、Session 或拦截器鉴权实现。正式验收文档中可以写“基于角色的访问控制设计”，但实现说明中应说明当前版本的权限控制仍需完善。"
    )


def section_modules(doc):
    add_heading(doc, "4. 功能模块设计", 1)
    add_callout(
        doc,
        "模块划分结论",
        "建议文档和小组分工均按业务模块划分，而不是按 Controller、Service、Mapper 文件夹划分。当前项目核心模块应包括用户权限、院系管理员、师生管理、答辩小组、评分统计、模板文档、签名日期和 AI 评语。"
    )

    modules = [
        ("用户登录与密码管理模块", "LoginController、UserController", "登录、带年份登录、获取教师答辩年份、退出、修改密码。", "user、tea_group_rel、tea_stu_rel"),
        ("院系与管理员管理模块", "InstituteController、AdminController", "院系增删改查，院系管理员创建、列表、修改、删除、重置密码。", "institute、user、user_inst_rel"),
        ("教师管理模块", "TeacherController", "教师列表、详情、创建、删除、加入/移出小组、设置组长、指导学生关系。", "user、tea_group_rel、tea_stu_rel、dbgroup"),
        ("学生与选题管理模块", "StudentController", "学生分页列表、未分配学生、学生增删改查、分配小组、设置论文/设计题目。", "student、dbinfo、tea_stu_rel、group_defense"),
        ("答辩年份与小组管理模块", "DefenseController、GroupController", "答辩年份创建/删除/统计，小组创建/修改/删除，成员查询与移除。", "dbgroup、dbinfo、tea_group_rel"),
        ("评分与成绩统计模块", "DefenseController、DefenseServiceImpl", "教师评分明细保存、论文/设计评分项区分、平均分计算、大组成绩、调节系数。", "dbinfo、group_defense、dbgroup"),
        ("模板与文档管理模块", "TemplateController、TemplateServiceImpl", "Word 模板上传、下载、删除、占位符校验、模板列表、日期应用。", "template、placeholder_config、date_config"),
        ("签名与日期配置模块", "SignatureController、DateConfigController", "签名上传与查询，答辩日期、评阅日期保存和读取。", "user、date_config"),
        ("AI 评语生成模块", "AiController", "AI 参数配置、提示词保存/读取、答辩评语生成和保存。", "config/*.json、dbinfo"),
        ("统计看板模块", "InstAdminStatsController", "院系管理员首页统计，统计学生、教师、年份、小组等数量。", "student、user、dbgroup、tea_group_rel"),
    ]
    add_table(doc, ["模块", "主要代码", "功能说明", "关联数据"], modules, [1900, 2300, 3600, 1560], font_size=8.8)

    add_heading(doc, "4.1 用户登录与权限模块", 2)
    add_body(doc, "该模块负责系统入口认证，支持普通登录和教师按年份登录。登录成功后返回用户信息、角色、院系和年份相关信息，为前端路由和功能入口控制提供依据。")
    add_bullets(
        doc,
        [
            "普通登录：校验用户名与密码，识别超级管理员、院系管理员、教师角色。",
            "带年份登录：教师选择答辩年份后登录，系统验证教师是否具备该年份相关权限。",
            "教师年份查询：返回教师可参与的答辩年份列表。",
            "密码修改：通过用户模块修改当前用户密码。",
        ],
    )

    add_heading(doc, "4.2 院系与管理员管理模块", 2)
    add_body(doc, "该模块用于维护学校或学院组织结构，并为院系绑定管理员账号。创建院系管理员时，系统会校验院系是否存在、是否已有管理员，然后创建用户并建立院系关联。")
    add_bullets(
        doc,
        [
            "院系维护：新增、修改、删除、查询院系列表和数量。",
            "管理员维护：创建、修改、删除院系管理员，重置默认密码。",
            "院系绑定：一个院系可绑定一个管理员，删除管理员前需清理院系关联。",
        ],
    )

    add_heading(doc, "4.3 教师管理模块", 2)
    add_body(doc, "教师模块负责维护教师基础信息和教师在答辩流程中的关系数据。教师既可以作为指导教师，也可以加入答辩小组，部分教师可被设置为答辩组长。")
    add_bullets(
        doc,
        [
            "教师账号创建时写入 user 表，并设置 role=2。",
            "教师加入小组时写入 tea_group_rel 表，同一年原则上只能参加一个小组。",
            "设置组长时检查小组是否已有组长，避免一个小组出现多个组长。",
            "指导关系通过 tea_stu_rel 表维护，用于教师查看指导学生和教师年份权限判断。",
        ],
    )

    add_heading(doc, "4.4 学生与选题管理模块", 2)
    add_body(doc, "学生模块负责学生基础信息、院系归属、指导教师关系和答辩分组。学生被分配小组后，dbinfo 表保存答辩信息、题目、类型和成绩相关字段。")
    add_bullets(
        doc,
        [
            "支持按院系分页查询学生列表。",
            "支持查询未分配答辩小组的学生。",
            "支持学生创建、修改、删除和按学号搜索。",
            "支持为学生分配答辩小组，并设置论文或设计类型。",
            "支持设置论文/设计题目，为后续评分和模板生成提供数据。"
        ],
    )

    add_heading(doc, "4.5 答辩年份与小组管理模块", 2)
    add_body(doc, "答辩年份是系统组织数据的主维度。院系管理员可以按年份创建答辩小组，设置组长、容量和成员，并将学生分配到对应小组。")
    add_bullets(
        doc,
        [
            "年份管理：新增、删除、查询全部年份及各年份统计。",
            "小组管理：创建、更新、删除答辩小组，维护组长与最大学生数。",
            "成员管理：查询小组学生列表，移除小组学生。",
            "院系统计：按年份和院系统计小组数量与答辩学生数量。",
        ],
    )

    add_heading(doc, "4.6 评分与成绩统计模块", 2)
    add_body(doc, "评分模块是本项目的核心业务。教师评分时，系统根据学生考核类型区分毕业论文和毕业设计评分项，保存每位教师评分明细，并将多位教师总分求平均后写入 total_score。")
    add_table(
        doc,
        ["考核类型", "评分项", "存储方式"],
        [
            ("毕业论文", "论文质量、汇报表现、答问情况、总分", "保存到 teacher_scores JSON，并更新 dbinfo.total_score。"),
            ("毕业设计", "设计质量1/2/3、设计汇报、答问1/2、总分", "保存到 teacher_scores JSON，并更新 dbinfo.total_score。"),
            ("大组评分", "各小组第一名的大组成绩", "保存到 group_defense.major_score。"),
            ("调节系数", "major_score / group_score", "按小组保存 adjustment_coefficient，用于成绩平衡。"),
        ],
        [1600, 4000, 3760],
        font_size=9,
    )

    add_heading(doc, "4.7 模板与文档管理模块", 2)
    add_body(doc, "模板模块负责管理答辩相关 Word 表格模板。系统支持上传 doc/docx 文件、提取占位符、校验必填占位符、下载模板、删除模板和配置日期。")
    add_bullets(
        doc,
        [
            "template 表保存模板名称、类型、文件路径、文件名、大小、更新时间和更新人。",
            "placeholder_config 表保存不同模板类型所需占位符，例如学生姓名、学号、题目、成绩、签名等。",
            "date_config 表保存答辩日期和评阅日期，可被模板生成逻辑复用。",
            "当前 TemplateServiceImpl 中 uploadTemplate 调用了 validateTemplatePlaceholders，但该私有方法当前直接返回 true；validateTemplate 接口具备实际占位符检查逻辑。"
        ],
    )

    add_heading(doc, "4.8 签名、日期与 AI 评语模块", 2)
    add_body(doc, "签名模块用于上传教师或管理员签名图片，日期模块用于统一配置答辩日期和评阅日期，AI 模块用于读取提示词和模型配置并调用 DashScope 生成答辩评语。")
    add_bullets(
        doc,
        [
            "签名文件限制为 png、jpg，默认存放在 uploads/signatures。",
            "AI 配置包含 apiKey、modelName 等参数，保存在 resources/config/ai.json。",
            "论文和设计分别使用 thesisprompt.json、designprompt.json 作为默认提示词。",
            "生成后的评语可保存到学生答辩信息中，支撑后续模板归档。",
        ],
    )


def section_database(doc):
    add_heading(doc, "5. 数据库设计", 1)
    add_body(doc, "数据库初始化脚本位于 src/main/resources/sql/init.sql。核心数据表围绕用户、院系、学生、答辩小组、评分、模板和配置展开。")
    add_table(
        doc,
        ["表名", "用途", "关键字段"],
        [
            ("user", "用户基础表，存储管理员、院系管理员、教师账号", "id、pwd、role、real_name、phone、email、signature"),
            ("institute", "院系表", "id、name、user_id"),
            ("user_inst_rel", "用户与院系关联表", "user_id、inst_id"),
            ("student", "学生基础表", "id、name、phone、email、institute_id"),
            ("tea_stu_rel", "教师指导学生关系", "tea_id、stu_id、year"),
            ("dbgroup", "答辩小组表", "id、year、admin_id、max_student_count、adjustment_coefficient"),
            ("tea_group_rel", "教师与答辩小组关系", "teacher_id、group_id、is_leader"),
            ("dbinfo", "学生答辩信息与成绩表", "gid、stu_id、type、title、reviewer_id、total_score、comment、advisor_id、teacher_scores"),
            ("group_defense", "大组成绩表", "group_id、stu_id、major_score"),
            ("template", "Word 模板表", "id、name、type、file_path、file_name、file_size、updated_by、updated_at"),
            ("placeholder_config", "模板占位符配置表", "template_type、placeholder_key、description、is_required"),
            ("date_config", "日期配置表", "config_key、config_value"),
        ],
        [1700, 3400, 4260],
        font_size=8.5,
    )

    add_heading(doc, "5.1 主要表关系", 2)
    add_bullets(
        doc,
        [
            "student.institute_id 关联 institute.id，表示学生所属院系。",
            "user_inst_rel 关联 user 与 institute，表示用户管理或隶属的院系。",
            "tea_stu_rel 关联教师 user.id 与 student.id，表示指导关系及年份。",
            "tea_group_rel 关联教师 user.id 与 dbgroup.id，表示教师加入答辩小组及是否为组长。",
            "dbinfo 通过 stu_id 和 gid 关联学生与答辩小组，并保存题目、考核类型、成绩、评语和教师评分 JSON。",
            "group_defense 使用 group_id 与 stu_id 记录大组成绩。",
            "template 与 placeholder_config 通过模板类型 type / template_type 关联。",
        ],
    )

    add_heading(doc, "5.2 成绩数据设计", 2)
    add_body(doc, "dbinfo.teacher_scores 使用 JSON 保存多位教师评分明细。该设计便于在一个字段中保存动态评分项，但统计查询和结构校验会比规范化表结构更复杂。")
    add_kv_table(
        doc,
        [
            ("论文评分 JSON", "包含 teacher_id、teacher_name、paper_quality、presentation、qa_performance、total_score。"),
            ("设计评分 JSON", "包含 teacher_id、teacher_name、design_quality1/2/3、design_presentation、design_qa1/2、total_score。"),
            ("total_score", "由当前学生该小组下所有教师评分总分取平均值后更新。"),
            ("调节系数", "按每个小组第一名的大组成绩和小组成绩计算，保存到 dbgroup.adjustment_coefficient。"),
        ],
    )

    add_heading(doc, "5.3 模板数据设计", 2)
    add_body(doc, "模板表 template 管理 7 类答辩相关 Word 模板。placeholder_config 表为每类模板配置必填占位符，用于验证上传模板是否符合系统生成要求。")


def section_api(doc):
    add_heading(doc, "6. 接口设计", 1)
    add_body(doc, "本节按照业务模块归纳主要接口。接口统一返回 Result 或类似 JSON 结构，常见字段包括 code、message、data。部分 Controller 直接返回 ResponseEntity<Map>，建议后续统一响应格式。")

    api_rows = [
        ("登录", "POST", "/login", "普通账号登录"),
        ("登录", "POST", "/loginWithYear", "教师按答辩年份登录"),
        ("登录", "GET", "/defenseYears/{teacherId}", "查询教师可用答辩年份"),
        ("用户", "POST", "/user/changePassword", "修改密码"),
        ("院系", "GET", "/institute/list", "查询院系列表"),
        ("院系", "POST", "/institute/add", "新增院系"),
        ("院系", "POST", "/institute/update", "修改院系"),
        ("院系", "POST", "/institute/delete", "删除院系"),
        ("管理员", "POST", "/admin/create", "创建院系管理员"),
        ("管理员", "GET", "/admin/list", "查询院系管理员列表"),
        ("管理员", "POST", "/admin/resetPassword", "重置管理员密码"),
        ("教师", "GET", "/teachers/list", "查询教师列表"),
        ("教师", "POST", "/teachers/create", "创建教师"),
        ("教师", "POST", "/teachers/add-to-group", "添加教师到答辩小组"),
        ("教师", "POST", "/teachers/set-defense-leader", "设置答辩组长"),
        ("学生", "GET", "/students/list", "分页查询学生"),
        ("学生", "GET", "/students/listunassign", "查询未分配学生"),
        ("学生", "POST", "/students/create", "创建学生"),
        ("学生", "POST", "/students/assign-group", "分配答辩小组"),
        ("学生", "POST", "/students/settitle", "设置论文/设计题目"),
        ("答辩年份", "POST", "/defense/yearadd", "新增答辩年份"),
        ("答辩年份", "GET", "/defense/allyear", "查询年份与统计信息"),
        ("答辩评分", "POST", "/defense/save-score", "保存教师评分"),
        ("答辩评分", "POST", "/defense/save-major-score", "保存大组成绩"),
        ("答辩评分", "POST", "/defense/save-coefficients", "计算并保存调节系数"),
        ("小组", "GET", "/groups/all", "查询答辩小组"),
        ("小组", "POST", "/groups/update", "创建或更新小组"),
        ("小组", "GET", "/groups/studentlist", "查询小组学生"),
        ("模板", "POST", "/templates/upload", "上传 Word 模板"),
        ("模板", "POST", "/templates/validate", "校验模板占位符"),
        ("模板", "GET", "/templates", "查询模板列表"),
        ("模板", "GET", "/templates/download/{id}", "下载模板"),
        ("日期", "POST", "/date-config/save", "保存答辩/评阅日期"),
        ("签名", "POST", "/signature/upload", "上传签名图片"),
        ("AI", "POST", "/ai/comment/generate", "生成 AI 答辩评语"),
        ("AI", "POST", "/ai/comment/save", "保存 AI 评语"),
    ]
    add_table(doc, ["模块", "方法", "路径", "说明"], api_rows, [1400, 1000, 3300, 3660], font_size=8.2)

    add_heading(doc, "6.1 响应格式建议", 2)
    add_code(doc, '{ "code": 200, "message": "success", "data": { ... } }')
    add_body(doc, "建议后续将所有接口统一为 Result<T> 响应，避免部分接口使用 success/code/message，部分接口使用 code/message/data，降低前端适配成本。")


def section_flows(doc):
    add_heading(doc, "7. 关键业务流程", 1)
    add_heading(doc, "7.1 教师评分流程", 2)
    add_numbered(
        doc,
        [
            "教师选择答辩年份并登录系统。",
            "前端根据教师所在小组展示学生列表和评分表。",
            "教师提交评分数据，接口 /defense/save-score 接收 stu_id、teacher_id、group_id、type 和各评分项。",
            "后端根据 type 构建论文或设计评分对象。",
            "系统读取 dbinfo.teacher_scores 中已有评分 JSON，更新当前教师评分或追加新评分。",
            "系统计算所有教师 total_score 的平均值，写入 dbinfo.total_score。",
            "评分明细 JSON 和平均总分共同保存，供统计和模板生成使用。",
        ],
    )

    add_heading(doc, "7.2 调节系数计算流程", 2)
    add_numbered(
        doc,
        [
            "系统查询指定年份每个答辩小组的第一名学生。",
            "管理员为第一名学生录入大组成绩 major_score。",
            "系统读取该学生所在小组成绩 group_score。",
            "调节系数按 major_score / group_score 计算并保留三位小数。",
            "调节系数保存到对应答辩小组，用于后续成绩平衡。"
        ],
    )

    add_heading(doc, "7.3 模板上传与校验流程", 2)
    add_numbered(
        doc,
        [
            "管理员选择模板类型并上传 doc/docx 文件。",
            "系统检查文件是否为空、扩展名是否合法、大小是否超过限制。",
            "系统读取模板类型对应的必填占位符。",
            "系统解析 Word 文档段落和表格，提取 {{xxx}} 或 ${xxx} 格式占位符。",
            "若缺少必填占位符，返回缺失列表；若通过校验，则保存模板文件并更新 template 表。",
        ],
    )

    add_heading(doc, "7.4 AI 评语生成流程", 2)
    add_numbered(
        doc,
        [
            "系统读取论文或设计类型对应的提示词模板。",
            "前端将学生姓名、题目、指导教师、答辩成绩等信息拼接进提示词。",
            "后端读取 ai.json 中的 apiKey 和 modelName。",
            "后端调用 DashScope Generation API 生成评语。",
            "用户确认后调用 /ai/comment/save 保存评语。",
        ],
    )


def section_deploy(doc):
    add_heading(doc, "8. 运行部署说明", 1)
    add_heading(doc, "8.1 环境要求", 2)
    add_table(
        doc,
        ["环境", "建议版本", "说明"],
        [
            ("JDK", "17", "pom.xml 中 java.version 为 17。"),
            ("Maven", "使用项目自带 mvnw/mvnw.cmd", "避免本机 Maven 版本差异。"),
            ("MySQL", "8.x 或兼容版本", "创建 manager 数据库并执行初始化脚本。"),
            ("Node.js", "可选", "仅在需要运行 server.js 辅助模板服务时使用。"),
        ],
        [1600, 2500, 5260],
    )

    add_heading(doc, "8.2 数据库初始化", 2)
    add_code(doc, "CREATE DATABASE manager DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;")
    add_code(doc, "mysql -u root -p manager < src/main/resources/sql/init.sql")
    add_body(doc, "初始化完成后，需检查 application.properties 中的数据库地址、用户名和密码是否与本机一致。")

    add_heading(doc, "8.3 启动 Spring Boot 后端", 2)
    add_code(doc, "mvnw.cmd spring-boot:run")
    add_body(doc, "服务启动后默认访问地址为 http://localhost:8081。前端调用接口时应以该地址作为后端 baseURL。")

    add_heading(doc, "8.4 可选 Node.js 服务", 2)
    add_body(doc, "server.js 中存在一个 Express 模板管理服务，默认端口为 3000。当前 Java 后端已包含 /templates 接口，因此 Node 服务更适合作为历史实现或调试参考。")
    add_code(doc, "npm install")
    add_code(doc, "npm run dev")

    add_heading(doc, "8.5 文件上传目录", 2)
    add_bullets(
        doc,
        [
            "模板文件默认保存到 ./uploads/templates/。",
            "签名图片默认保存到 ./uploads/signatures/。",
            "生产环境建议将上传目录移出代码目录，并配置备份策略。",
        ],
    )


def section_tests(doc):
    add_heading(doc, "9. 测试设计", 1)
    add_body(doc, "当前仓库未看到完整自动化测试用例，建议按业务模块设计接口测试、数据库一致性测试和文件上传测试。")
    test_rows = [
        ("T01", "登录", "输入正确账号密码", "返回用户信息、角色和权限数据"),
        ("T02", "教师年份登录", "教师选择无权限年份", "返回无权限或登录失败提示"),
        ("T03", "院系管理", "新增院系后查询列表", "列表包含新增院系"),
        ("T04", "管理员管理", "为已有管理员院系再次创建管理员", "返回该院系已有管理员"),
        ("T05", "教师入组", "同一年重复加入不同小组", "返回同一年只能参加一个小组"),
        ("T06", "学生分组", "将未分配学生加入小组", "dbinfo 生成或更新对应记录"),
        ("T07", "论文评分", "提交论文评分项", "teacher_scores JSON 更新，total_score 为平均分"),
        ("T08", "设计评分", "提交设计评分项", "设计评分字段保存，论文字段为空"),
        ("T09", "大组成绩", "录入 0-100 范围外成绩", "返回成绩范围错误"),
        ("T10", "模板校验", "上传缺少占位符的模板", "返回缺失占位符列表"),
        ("T11", "签名上传", "上传非 png/jpg 文件", "返回文件类型错误"),
        ("T12", "AI 评语", "配置有效 apiKey 和模型名", "成功返回模型生成结果"),
    ]
    add_table(doc, ["编号", "模块", "测试场景", "期望结果"], test_rows, [900, 1500, 3600, 3360], font_size=8.5)

    add_heading(doc, "9.1 验收重点", 2)
    add_bullets(
        doc,
        [
            "接口返回格式是否稳定，前端是否能统一解析。",
            "评分保存后，教师评分明细和平均分是否一致。",
            "删除院系、教师、小组、学生时，外键关系是否正确处理。",
            "模板上传后物理文件和数据库记录是否一致。",
            "AI 配置错误时是否有明确错误提示，避免后端异常直接暴露。",
        ],
    )


def section_team(doc):
    add_heading(doc, "10. 项目分工建议", 1)
    add_body(doc, "分工建议按业务模块划分，避免按 Controller、Service、Mapper 机械拆分。业务模块分工更符合答辩文档表达，也更容易说明每位成员完成的功能价值。")
    add_table(
        doc,
        ["成员", "建议负责模块", "主要工作内容"],
        [
            ("成员A", "用户登录、权限、院系与管理员管理", "实现登录、角色识别、院系维护、院系管理员创建、修改、删除和密码重置。"),
            ("成员B", "教师管理、学生管理、指导关系", "实现教师账号维护、学生增删改查、指导学生关系、学生题目维护。"),
            ("成员C", "答辩年份、小组、分组管理", "实现答辩年份维护、答辩小组创建、设置组长、教师入组、学生分组。"),
            ("成员D", "评分成绩、调节系数、统计", "实现论文/设计评分、大组成绩、调节系数、院系管理员统计看板。"),
            ("成员E", "模板、签名、日期、AI 评语", "实现 Word 模板上传校验、日期配置、签名上传、AI 提示词和评语生成。"),
        ],
        [1000, 2800, 5560],
        font_size=9,
    )

    add_heading(doc, "10.1 写进论文或报告时的模块名称", 2)
    add_bullets(
        doc,
        [
            "不要写“Controller 模块、Service 模块、Mapper 模块”，这属于技术层次，不适合作为项目分工模块。",
            "建议写“用户与权限模块、师生信息模块、答辩小组模块、成绩评分模块、模板文档模块、AI 评语模块”。",
            "如果成员较少，可以合并为三大块：基础信息管理、答辩业务管理、文档与智能辅助。"
        ],
    )


def section_risks(doc):
    add_heading(doc, "11. 风险问题与优化建议", 1)
    add_table(
        doc,
        ["问题", "影响", "建议"],
        [
            ("中文注释和部分 SQL 示例存在编码异常", "阅读和维护困难，可能影响文档展示", "统一使用 UTF-8 保存源码、SQL 和配置文件。"),
            ("application.properties 直接写数据库账号密码", "存在安全风险，不适合生产环境", "改为环境变量或外部配置文件。"),
            ("权限认证不完整", "接口可能被未授权调用", "增加 JWT、拦截器和角色权限校验。"),
            ("Java 与 Node 模板服务并存", "接口重复，部署复杂", "明确主实现，建议保留 Java 后端，Node 作为历史参考。"),
            ("响应格式不统一", "前端处理复杂，异常情况难统一", "统一 Result<T> 响应结构和错误码。"),
            ("模板上传校验逻辑不完全一致", "uploadTemplate 可能绕过实际占位符校验", "复用 validateTemplate 的占位符检查逻辑。"),
            ("评分明细使用 JSON 存储", "统计查询复杂，数据约束弱", "若后期统计复杂，可拆分 teacher_score 明细表。"),
            ("缺少自动化测试", "回归风险高", "补充 Controller 层接口测试和 Service 单元测试。"),
        ],
        [2300, 3000, 4060],
        font_size=8.5,
    )

    add_heading(doc, "11.1 后续可扩展方向", 2)
    add_bullets(
        doc,
        [
            "增加基于角色的接口权限控制，细化超级管理员、院系管理员、教师、答辩组长权限。",
            "增加答辩材料自动生成能力，将模板占位符与学生、成绩、签名数据自动合并。",
            "增加导入导出功能，例如批量导入学生、教师和导出成绩表。",
            "增加操作日志，记录管理员创建、删除、评分、模板上传等关键操作。",
            "增加接口文档平台，例如 Knife4j 或 Swagger UI，方便前后端联调。",
        ],
    )


def section_appendix(doc):
    add_heading(doc, "附录：代码依据与文件映射", 1)
    add_table(
        doc,
        ["文件/目录", "说明"],
        [
            ("pom.xml", "Spring Boot、MyBatis-Plus、DashScope、Apache POI、MySQL 等依赖配置。"),
            ("package.json / server.js", "Node.js Express 辅助模板管理服务。"),
            ("src/main/java/com/world/back/controller", "全部 REST Controller。"),
            ("src/main/java/com/world/back/serviceImpl", "主要业务逻辑实现。"),
            ("src/main/java/com/world/back/mapper", "MyBatis Mapper 数据访问层。"),
            ("src/main/java/com/world/back/entity", "实体类和响应类。"),
            ("src/main/resources/application.properties", "端口、数据库、上传目录等配置。"),
            ("src/main/resources/config/ai.json", "AI 模型配置。"),
            ("src/main/resources/config/thesisprompt.json", "毕业论文评语提示词。"),
            ("src/main/resources/config/designprompt.json", "毕业设计评语提示词。"),
            ("src/main/resources/sql/init.sql", "数据库建表和初始数据脚本。"),
        ],
        [3400, 5960],
        font_size=9,
    )


def finish_headers_footers(doc):
    for section in doc.sections:
        header = section.header
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.text = ""
        r = hp.add_run("毕业设计管理系统后端项目文档")
        set_run_font(r, size=9, color=MUTED)
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        add_page_number(fp)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    cover(doc)
    revision_and_toc(doc)
    section_project_overview(doc)
    section_architecture(doc)
    section_roles(doc)
    section_modules(doc)
    section_database(doc)
    section_api(doc)
    section_flows(doc)
    section_deploy(doc)
    section_tests(doc)
    section_team(doc)
    section_risks(doc)
    section_appendix(doc)
    finish_headers_footers(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE.resolve())


if __name__ == "__main__":
    build()
